from pathlib import Path

from docx import Document
from docx.enum.text import WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


OUT = Path("output/doc/Zhiming_Liu_中文更新版.docx")

BODY_FONT = "DengXian"
HEADING_FONT = "SimHei"
LABEL_FONT = "KaiTi"
LATIN_FONT = "Times New Roman"


def set_font(run, name=BODY_FONT, size=10.5, bold=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def paragraph(doc, text="", size=10.5, bold=False, font=BODY_FONT, before=0, after=0, line=1.0):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    r = p.add_run(text)
    set_font(r, font, size, bold)
    return p


def add_bottom_border(paragraph):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    pBdr.append(bottom)


def section(doc, title):
    p = paragraph(doc, title, size=15.5, bold=True, font=HEADING_FONT, before=6, after=2)
    add_bottom_border(p)
    return p


def tabbed(doc, left, mid=None, degree=None, right=None, size=10.5):
    p = paragraph(doc, "", size=size, after=0)
    stops = p.paragraph_format.tab_stops
    stops.add_tab_stop(Inches(2.15), WD_TAB_ALIGNMENT.LEFT)
    stops.add_tab_stop(Inches(4.15), WD_TAB_ALIGNMENT.LEFT)
    stops.add_tab_stop(Inches(7.0), WD_TAB_ALIGNMENT.RIGHT)
    parts = [left]
    if mid is not None:
        parts.append(mid)
    if degree is not None:
        parts.append(degree)
    if right is not None:
        parts.append(right)
    r = p.add_run("\t".join(parts))
    set_font(r, size=size, bold=True)
    return p


def bullet(doc, text, level=0):
    p = doc.add_paragraph(style=None)
    p.paragraph_format.left_indent = Inches(0.23 + level * 0.15)
    p.paragraph_format.first_line_indent = Inches(-0.12)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run("•   " + text)
    set_font(r, size=10.3)
    return p


def build():
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)

    sec = doc.sections[0]
    sec.top_margin = Inches(0.55)
    sec.bottom_margin = Inches(0.45)
    sec.left_margin = Inches(0.43)
    sec.right_margin = Inches(0.43)

    paragraph(doc, "刘志明", size=24, bold=True, font=HEADING_FONT, after=10)
    p = paragraph(doc, "", after=12)
    for text, bold, font in [
        ("电话：", True, BODY_FONT),
        ("+852 9631 0816", True, LATIN_FONT),
        ("     |     邮箱：", True, BODY_FONT),
        ("lzm200303@gmail.com", True, LATIN_FONT),
        ("     |     Github: ", True, LATIN_FONT),
        ("https://github.com/Lzm03", True, LATIN_FONT),
    ]:
        r = p.add_run(text)
        set_font(r, font, 11.5, bold)

    section(doc, "教育背景")
    tabbed(doc, "布里斯托大学(英国)", "计算机科学", "本硕连读", "2021.09 - 2025.07", 11)
    paragraph(doc, "成绩：英国 二等一学位，绩点：3.6/4.0")
    paragraph(doc, "核心课程：算法，数据驱动计算机科学，计算机系统，计算机图形学")
    tabbed(doc, "香港理工大学(中国香港)", "元宇宙科技", "硕士", "2025.09 - 2026.07", 11)
    paragraph(doc, "核心课程：机器学习，自然语言处理，计算机视觉")

    section(doc, "论文发表")
    for pub in [
        "[1] Z Liu, N Anantrasirichai, RMFAT: 循环多尺度特征大气湍流缓解模型，AAAI。",
        "[2] P Hill, Z Liu, A Achim, D Bull, N Anantrasirichai, DMAT: 大气湍流缓解与目标检测联合端到端框架，WACV。",
        "[3] Z Liu, P Hill, N Anantrasirichai, JDATT: 面向大气湍流缓解与目标检测的联合蒸馏框架，BMVC。",
        "[4] P. Hill, Z. Liu and N. Anantrasirichai, MAMAT: 基于 3D Mamba 的大气湍流去除及目标检测能力，AVSS 2025。",
    ]:
        paragraph(doc, pub, size=9.5)

    section(doc, "研究经历")
    tabbed(doc, "大气湍流缓解研究", None, None, "2025.05 - 至今", 11)
    paragraph(doc, "视觉信息实验室，布里斯托大学，英国布里斯托", bold=True)
    paragraph(doc, "研究内容:", font=LABEL_FONT, size=11)
    for item in [
        "聚焦基于深度学习的大气湍流缓解方法，以及畸变条件下的实时目标检测；研究内容包括循环网络、联合复原-检测框架和3D Mamba架构。",
        "开发 MAMAT，一套基于 3D Mamba 的湍流去除系统，具备良好的时序一致性并提升下游检测能力；发表于 AVSS 2025。",
        "设计 JDATT 联合蒸馏框架，将复原与检测知识融合，提升学生模型在湍流场景中的表现；以第一作者论文发表于 BMVC 2025。",
        "提出 DMAT 端到端框架，联合大气湍流缓解与目标检测，在严重畸变下显著提升检测鲁棒性；发表于 WACV 2026。",
        "构建 RMFAT 循环多尺度特征湍流缓解模型，针对模糊、几何畸变、抖动和时序不一致等问题进行建模；以第一作者论文发表于 AAAI 2026。",
    ]:
        bullet(doc, item)

    section(doc, "实习经历")
    jobs = [
        ("视觉信息实验室", "布里斯托大学，英国布里斯托", "2025.05 - 2025.12", "研究助理", [
            "模型复现：复现基线模型，准备数据集，并搭建基准测试流程。",
            "实验执行：开展受控实验，并维护具备可复现配置的训练流程。",
            "研究支持：协助进行结果分析、模型优化和研究成果整理。",
        ]),
        ("Index Academy", "中国香港，九龙", "2025.09 - 至今", "全栈开发工程师 - AI教育平台", [
            "平台开发：为中小学开发 AI 驱动的学习平台。",
            "功能实现：实现全栈功能，包括 3D 角色创建、语音定制和课程内容集成。",
            "系统优化：提升系统稳定性与交互性能，支持实时、可用于课堂的 AI 教学体验。",
        ]),
        ("深圳市维佳科创科技有限公司", "中国深圳", "2024.06 - 2024.08", "游戏开发实习生", [
            "使用 Unity 设计面向幼儿的互动教育游戏。",
            "构建基于物理的交互、拖拽逻辑和实时反馈机制。",
            "使用 C# 编写场景交互逻辑和游戏玩法行为。",
        ]),
    ]
    for company, loc, date, role, items in jobs:
        tabbed(doc, company, loc, None, date, 10.6)
        paragraph(doc, role, font=LABEL_FONT, size=10.8)
        for item in items:
            bullet(doc, item)

    doc.add_page_break()
    section(doc, "项目经历")
    projects = [
        ("校内项目", "大气湍流图像恢复与目标检测联合优化", "2024.12 - 2025.04",
         "项目描述：大气湍流导致的长距离图像/视频失真（如波纹、模糊）严重影响目标识别与检测的准确性。现有方法（如 CNN、Transformer、3D Mamba 架构）虽能提升图像质量，但依赖复杂网络和大参数量，难以满足实时性需求。同时，主流检测模型（如 Faster R-CNN、YOLO）在湍流环境下计算开销大，导致实时性下降。",
         ["筛选出湍流数据集上 PSNR、SSIM 最优的恢复模型及 mAP 最优的检测模型。", "构建端到端联合训练系统，恢复质量（PSNR/SSIM）与检测精度（mAP）均接近或优于基线模型。", "通过对比实验验证联合框架优势：模型体积减少 10%，推理速度提升 30%，适用于边缘设备部署。"]),
        ("校内项目", "图像超分辨率方法对比研究", "2024.10 - 2024.11",
         "项目描述：研究基于深度学习的图像超分辨率技术，对比 DIP、GAN 和 Deep Unfolding 在噪声和降分辨率下的性能。使用 DIV2K 数据集，评估细节恢复(LPIPS)、结构相似性(SSIM)和视觉质量(PSNR)。",
         ["DIP 在细节还原表现最优（PSNR 31.60/SSIM 0.9387）。", "GAN 生成图像感知质量最佳（LPIPS 0.3284）。", "Deep Unfolding 在噪声/降采样场景鲁棒性最强（PSNR 34.85/SSIM 0.9532）。"]),
        ("校内项目（6人小组）", "气味陷阱游戏", "2024.01 - 2024.05",
         "项目描述：气味陷阱是一个多人在线游戏，背景设定在一个荒废的大学里，玩家可以扮演人类或变异芝士角色进行对战。游戏支持最多 20 人同时在线，强调策略和团队协作。",
         ["开发基于 Unity3D 和 C# 的多人在线对战游戏，支持 20 人同时在线。", "负责核心游戏逻辑：角色控制、技能系统和阵营交互逻辑。", "实现 Photon 引擎网络同步，优化匹配与实时对战体验，确保低延迟与稳定性。", "设计策略地图与气味可视化系统，通过 Unity 粒子系统动态模拟气味扩散。", "开发 UI 交互界面，实时显示体力条、技能冷却等关键信息，优化玩家操作体验。"]),
    ]
    for kind, title, date, desc, items in projects:
        tabbed(doc, kind, title, None, date, 10.6)
        paragraph(doc, desc, size=10.2)
        paragraph(doc, "项目成果:" if "项目描述" in desc and title != "气味陷阱游戏" else "主要工作:", font=LABEL_FONT, size=10.8)
        for item in items:
            bullet(doc, item)

    section(doc, "个人技能与荣誉")
    for item in [
        "外语能力：普通话（母语），英语（熟练），粤语（流利）",
        "编程技能：Python（numpy、pandas、scikit-learn、matplotlib、pytorch），C++，C#，Java，GoLand",
        "软件与工具：LaTeX，Markdown，Excel，Adobe Photoshop",
    ]:
        bullet(doc, item)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
