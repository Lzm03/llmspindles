from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


SRC = Path("/Users/liuzhiming/Desktop/cv/CV_RA.docx")
OUT = Path("output/doc/CV_RA_中文.docx")


TRANSLATIONS = {
    0: "刘志明（Jayden）",
    1: "电话：+852 9631 0816 | 邮箱：lzm200303@gmail.com | GitHub：https://github.com/Lzm03 ",
    2: "教育经历",
    3: "布里斯托大学，英国布里斯托\t\t\t\t\t\t\t\t\t\t     2021年9月 - 2025年7月",
    4: "计算机科学本硕连读工程硕士 | GPA：3.6/4.00",
    5: "核心课程：算法、数据驱动计算机科学、计算机系统、计算机图形学",
    6: "香港理工大学，香港九龙\t\t\t\t\t\t\t\t\t 2025年9月 - 2026年7月",
    7: "元宇宙科技理学硕士",
    8: "核心课程：机器学习、自然语言处理、计算机视觉",
    10: "论文发表",
    11: "[1] Z Liu, N Anantrasirichai, RMFAT：循环多尺度特征大气湍流缓解模型。AAAI人工智能会议论文集（AAAI）。",
    12: "[2] P Hill, Z Liu, A Achim, D Bull, N Anantrasirichai, DMAT：面向大气湍流缓解与目标检测联合任务的端到端框架。IEEE/CVF冬季计算机视觉应用会议（WACV）。",
    13: "[3] Z Liu, P Hill, N Anantrasirichai, JDATT：面向大气湍流缓解与目标检测的联合蒸馏框架，第36届英国机器视觉会议（BMVC）。",
    14: "[4] P. Hill, Z. Liu and N. Anantrasirichai, MAMAT：基于3D Mamba的大气湍流去除方法及其目标检测能力，2025 IEEE先进视觉与信号系统国际会议（AVSS）。",
    16: "研究经历",
    17: "大气湍流缓解研究                                        2025年5月 - 至今",
    18: "视觉信息实验室，布里斯托大学，英国布里斯托",
    19: "聚焦基于深度学习的大气湍流缓解方法，以及畸变条件下的实时目标检测；研究内容包括循环网络、联合复原-检测框架和3D Mamba架构。",
    20: "开发MAMAT，一套基于3D Mamba的湍流去除系统，具备良好的时序一致性并提升下游检测能力；发表于AVSS 2025。",
    21: "设计JDATT联合蒸馏框架，将复原与检测知识融合，以提升学生模型在湍流场景中的表现；以第一作者论文发表于BMVC 2025。",
    22: "提出DMAT端到端框架，用于联合大气湍流缓解与目标检测，在严重畸变下显著提升检测鲁棒性；发表于WACV 2026。",
    23: "构建RMFAT循环多尺度特征湍流缓解模型，针对模糊、几何畸变、抖动和时序不一致等问题进行建模；以第一作者论文发表于AAAI 2026。",
    25: "工作经历",
    26: "视觉信息实验室                                            布里斯托大学，英国布里斯托",
    27: "研究助理                                                                 2025年5月 - 2025年12月",
    28: "模型复现：复现基线模型，准备数据集，并搭建基准测试流程。",
    29: "实验执行：开展受控实验，并维护具备可复现配置的训练流程。",
    30: "研究支持：协助进行结果分析、模型优化和研究成果整理。",
    32: "Index Academy                                                                  香港九龙",
    33: "全栈开发工程师 - AI教育平台                                          2025年9月 - 至今",
    34: "平台开发：为中小学开发AI驱动的学习平台。",
    35: "功能实现：实现全栈功能，包括3D角色创建、语音定制和课程内容集成。",
    36: "系统优化：提升系统稳定性与交互性能，支持实时、可用于课堂的AI教学体验。",
    38: "深圳微加科创科技                                            中国深圳",
    39: "游戏开发实习生                                                       2024年6月 - 2024年8月",
    40: "玩法设计：使用Unity为低龄儿童设计互动教育游戏。",
    41: "功能实现：构建基于物理的交互、拖拽逻辑和实时反馈机制。",
    42: "C#开发：使用C#编写场景交互逻辑和游戏玩法行为。",
    44: "技能",
    45: "语言：普通话（母语）、英语（熟练）、粤语（流利）",
    46: "编程语言：Python（numpy、pandas、scikit-learn、matplotlib、pytorch）、C++、C#、Java、GoLand",
    47: "软件与工具：LaTeX、Markdown、Excel、Adobe Photoshop",
}


def set_paragraph_text(paragraph, text):
    p = paragraph._p
    ppr = p.pPr
    for child in list(p):
        if child is not ppr:
            p.remove(child)
    run = paragraph.add_run(text)
    rpr = run._r.get_or_add_rPr()
    fonts = rpr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.append(fonts)
    fonts.set(qn("w:eastAsia"), "Songti SC")


def main():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document(SRC)
    for idx, text in TRANSLATIONS.items():
        set_paragraph_text(doc.paragraphs[idx], text)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
